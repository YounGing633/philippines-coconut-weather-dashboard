from __future__ import annotations
from pathlib import Path
import base64, html
from typing import Any
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter

THEME='#007370'


def _safe_sheet_name(name: str) -> str:
    return str(name)[:31].replace('/', '_').replace('\\','_')


def _write_df(ws, df: pd.DataFrame, max_rows: int | None = None):
    if df is None or df.empty:
        ws.cell(1,1,'无数据 / No data')
        return
    d=df.copy()
    if max_rows is not None and len(d)>max_rows:
        d=d.head(max_rows)
        ws.cell(1,1,f'仅展示前{max_rows}行；完整数据见data/history目录。')
        start=3
    else:
        start=1
    for cidx, col in enumerate(d.columns, 1):
        cell=ws.cell(start, cidx, col)
        cell.font=Font(bold=True, color='FFFFFF')
        cell.fill=PatternFill('solid', fgColor='007370')
        cell.alignment=Alignment(horizontal='center', vertical='center', wrap_text=True)
    for ridx, row in enumerate(d.itertuples(index=False), start+1):
        for cidx, val in enumerate(row, 1):
            if pd.isna(val): val=''
            cell=ws.cell(ridx, cidx, val)
            cell.alignment=Alignment(vertical='top', wrap_text=True)
            cell.border=Border(bottom=Side(style='hair', color='DDDDDD'))
    ws.freeze_panes=ws.cell(start+1,1).coordinate
    for col_idx in range(1, min(len(d.columns), 40)+1):
        letter=get_column_letter(col_idx)
        vals=[str(ws.cell(r, col_idx).value or '') for r in range(1, min(ws.max_row, 200)+1)]
        width=min(max(max([len(v) for v in vals]+[8])+2, 10), 48)
        ws.column_dimensions[letter].width=width


def _annotations_sheet(wb, sections: list[dict[str,Any]]):
    ws=wb.create_sheet('中文释义')
    ws['A1']='中文释义 / 读图说明'
    ws['A1'].font=Font(size=18, bold=True, color='007370')
    row=3
    for sec in sections:
        ws.cell(row,1,sec.get('title','说明')).font=Font(size=13,bold=True,color='007370')
        row+=1
        for b in sec.get('bullets',[]):
            ws.cell(row,1,'•')
            ws.cell(row,2,str(b))
            ws.cell(row,2).alignment=Alignment(wrap_text=True, vertical='top')
            row+=1
        if sec.get('comment'):
            ws.cell(row,2,sec['comment'])
            ws.cell(row,2).fill=PatternFill('solid', fgColor='FFF8E5')
            ws.cell(row,2).alignment=Alignment(wrap_text=True, vertical='top')
            row+=1
        row+=1
    ws.column_dimensions['A'].width=8; ws.column_dimensions['B'].width=110
    for i in range(1,row+1): ws.row_dimensions[i].height=24


def write_excel(out_path: str | Path, summary: dict, sections: list[dict[str,Any]], metrics: pd.DataFrame, group_summary: pd.DataFrame, forecast_metrics: pd.DataFrame, pagasa: pd.DataFrame, official: pd.DataFrame, locations: pd.DataFrame, quality: pd.DataFrame, history_recent: pd.DataFrame, source_log: pd.DataFrame, chart_paths: dict[str,Path], history_path: Path) -> Path:
    out=Path(out_path); out.parent.mkdir(parents=True, exist_ok=True)
    wb=Workbook(); ws=wb.active; ws.title='0_总览'
    ws['A1']='菲律宾椰子产区天气监测 / CNO-Copra Weather Monitor'
    ws['A1'].font=Font(size=18,bold=True,color='007370')
    ws['A3']='一句话结论'; ws['A3'].font=Font(size=13,bold=True)
    ws['B3']=summary.get('one_liner',''); ws['B3'].alignment=Alignment(wrap_text=True,vertical='top')
    ws.merge_cells('B3:J6')
    ws['A8']='详细解释'; ws['A8'].font=Font(size=13,bold=True)
    ws['B8']=summary.get('detail',''); ws['B8'].alignment=Alignment(wrap_text=True,vertical='top')
    ws.merge_cells('B8:J18')
    ws['A20']='完整历史日度数据'; ws['A20'].font=Font(size=13,bold=True)
    ws['B20']=str(history_path); ws['B20'].alignment=Alignment(wrap_text=True)
    ws.merge_cells('B20:J20')
    row=22; col=1
    for name, path in chart_paths.items():
        if Path(path).exists():
            ws.cell(row, col, name).font=Font(bold=True,color='007370')
            img=XLImage(str(path)); oldw,oldh=img.width,img.height
            neww=min(720, oldw); ratio=neww/max(oldw,1); img.width=neww; img.height=int(oldh*ratio)
            ws.add_image(img, f'{get_column_letter(col)}{row+1}')
            if col==1:
                col=9
            else:
                col=1; row+=28
    ws.column_dimensions['A'].width=18; ws.column_dimensions['B'].width=26
    ws.row_dimensions[3].height=70; ws.row_dimensions[8].height=150
    _annotations_sheet(wb, sections)
    sheet_data=[
        ('1_窗口指标', metrics),
        ('2_区域汇总', group_summary),
        ('3_未来预测_弥补', forecast_metrics),
        ('4_PAGASA式干旱', pagasa),
        ('5_PAGASA官方观点', official),
        ('6_产区权重', locations),
        ('7_数据质量', quality),
        ('8_近期日度数据', history_recent),
        ('9_数据源日志', source_log),
        ('10_指标说明', pd.DataFrame([
            {'指标':'rain_pct_normal','中文释义':'过去窗口实际累计降雨 / 1981-2020同日历日常年累计降雨。空值和-999不计入。'},
            {'指标':'risk_label','中文释义':'市场监测干旱等级，基于7/14/30/90天降雨距平、高温天数和有效数据覆盖率。'},
            {'指标':'pagasa_like_assessment','中文释义':'用PAGASA官方dry condition/dry spell/drought规则套用NASA降雨数据的近似估算，不替代官方发布。'},
            {'指标':'forecast_rain_to_30d_deficit_pct','中文释义':'未来预测降雨 / 过去30天降雨亏缺，用来判断是否能弥补。'},
            {'指标':'effective_observation_end_date','中文释义':'剔除空值/-999后，本次实际可用观测截止日。'},
        ])),
    ]
    for name, df in sheet_data:
        ws2=wb.create_sheet(_safe_sheet_name(name))
        _write_df(ws2, df, max_rows=20000 if name=='8_近期日度数据' else None)
    wb.save(out); return out


def _img64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode('ascii') if Path(path).exists() else ''


def _sections_html(sections):
    blocks=[]
    for sec in sections:
        lis=''.join(f'<li>{html.escape(str(x))}</li>' for x in sec.get('bullets',[]))
        comment=f"<p class='note'>{html.escape(sec.get('comment',''))}</p>" if sec.get('comment') else ''
        blocks.append(f"<div class='subcard'><h3>{html.escape(sec.get('title','说明'))}</h3><ul>{lis}</ul>{comment}</div>")
    return ''.join(blocks)


def write_html(out_path: str | Path, summary: dict, sections: list[dict], metrics: pd.DataFrame, group_summary: pd.DataFrame, forecast_metrics: pd.DataFrame, pagasa: pd.DataFrame, official: pd.DataFrame, quality: pd.DataFrame, chart_paths: dict[str,Path], history_path: Path) -> Path:
    out=Path(out_path); out.parent.mkdir(parents=True, exist_ok=True)
    css="""
    body{font-family:Arial,'Microsoft YaHei',sans-serif;margin:28px;background:#fff;color:#222}.title{font-size:26px;font-weight:700;color:#007370}.small{font-size:12px;color:#666}.card{border:1px solid #d9e6e3;border-radius:12px;padding:16px;margin:14px 0;box-shadow:0 2px 8px rgba(0,0,0,.05)}.subcard{background:#f7fbfa;border-left:5px solid #007370;border-radius:8px;padding:12px 14px;margin:12px 0}.subcard h3{margin:0 0 8px 0;color:#007370}.subcard li{margin:6px 0;line-height:1.55}.note{background:#fff8e5;border-left:4px solid #d59b00;padding:8px 10px;border-radius:6px;color:#555}table{border-collapse:collapse;width:100%;font-size:12px}th{background:#007370;color:white}td,th{border:1px solid #ddd;padding:6px;vertical-align:top}img{max-width:100%;height:auto}pre{white-space:pre-wrap;font-family:inherit;line-height:1.5}.grid{display:grid;grid-template-columns:1fr 1fr;gap:14px}@media(max-width:1200px){.grid{display:block}}
    """
    imgs=''.join([f"<div class='card'><h3>{html.escape(k)}</h3><img src='data:image/png;base64,{_img64(Path(p))}'/></div>" for k,p in chart_paths.items() if Path(p).exists()])
    tables=[]
    if not group_summary.empty: tables.append('<h3>区域汇总 / Regional weighted summary</h3>'+group_summary.head(30).to_html(index=False, border=0))
    if not metrics.empty:
        cols=[c for c in ['map_label','window_days','rain_pct_normal','rain_deficit_mm','temp_anom_c','risk_label','trend_label','risk_note'] if c in metrics.columns]
        tables.append('<h3>主产区窗口指标 / Major producing areas</h3>'+metrics[metrics['window_days'].isin([30,90])][cols].head(80).to_html(index=False, border=0))
    if not forecast_metrics.empty:
        cols=[c for c in ['map_label','forecast_window_days','forecast_rain_mm','forecast_rain_to_30d_deficit_pct','relief_label','base_30d_risk_label'] if c in forecast_metrics.columns]
        tables.append('<h3>未来预测与弥补判断 / Forecast relief</h3>'+forecast_metrics[forecast_metrics['forecast_window_days'].isin([7,16])][cols].head(80).to_html(index=False, border=0))
    if not pagasa.empty:
        cols=[c for c in ['map_label','pagasa_like_assessment','pagasa_like_zh','monthly_pct_latest3','coconut_mature_2025_share_pct'] if c in pagasa.columns]
        tables.append('<h3>PAGASA式干旱估算 / PAGASA-like assessment</h3>'+pagasa[cols].head(80).to_html(index=False, border=0))
    if not official.empty:
        cols=[c for c in ['title','url','english_core_text','中文核心观点'] if c in official.columns]
        tables.append('<h3>PAGASA官方观点 / Official views</h3>'+official[cols].head(20).to_html(index=False, escape=False, border=0))
    html_txt=f"""<!doctype html><html><head><meta charset='utf-8'><title>Philippines Coconut Weather Monitor</title><style>{css}</style></head><body>
    <div class='title'>菲律宾椰子产区天气监测 / CNO-Copra Weather Monitor</div>
    <div class='small'>Historical daily weather store + current risk + forecast relief + PAGASA official signals</div>
    <div class='card'><h2>一句话结论</h2><p>{html.escape(summary.get('one_liner',''))}</p></div>
    <div class='card'><h2>详细解释</h2><pre>{html.escape(summary.get('detail',''))}</pre></div>
    <div class='card'><h2>中文释义 / How to read this report</h2>{_sections_html(sections)}</div>
    <div class='grid'>{imgs}</div>
    <div class='card'><h2>完整历史日度数据</h2><p>{html.escape(str(history_path))}</p></div>
    <div class='card'>{''.join(tables)}</div>
    </body></html>"""
    out.write_text(html_txt, encoding='utf-8'); return out


def write_docx(out_path: str | Path, summary: dict, sections: list[dict], chart_paths: dict[str,Path]) -> Path | None:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception:
        print('[WARN] python-docx not installed; skip docx.')
        return None
    out=Path(out_path); out.parent.mkdir(parents=True, exist_ok=True)
    doc=Document()
    doc.add_heading('菲律宾椰子产区天气监测 / CNO-Copra Weather Monitor', level=1)
    doc.add_heading('一句话结论', level=2); doc.add_paragraph(summary.get('one_liner',''))
    doc.add_heading('详细解释', level=2)
    for line in summary.get('detail','').split('\n'):
        doc.add_paragraph(line)
    doc.add_heading('中文释义 / 如何看这份报告', level=2)
    for sec in sections:
        doc.add_heading(sec.get('title','说明'), level=3)
        for b in sec.get('bullets',[]): doc.add_paragraph(str(b), style='List Bullet')
        if sec.get('comment'): doc.add_paragraph(sec['comment'])
    for name, p in chart_paths.items():
        if Path(p).exists():
            doc.add_heading(name, level=2)
            doc.add_picture(str(p), width=Inches(6.5))
    doc.save(out); return out
